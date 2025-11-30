import os
import logging
import sqlite3
import tempfile
import requests
import json
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn

# ---------------- Настройки ----------------
BOT_TOKEN = "7388004511:AAEVOtZtcFfEae3hNeLe9ZlqWOY3ZcfH9bY"
DB_PATH = os.path.join(os.path.dirname(__file__), "audit.db")
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "data", "act_fundament_template.docx")

CHECKLIST = [
    "Марка бетона соответствует проекту ?",
    "Толщина защитного слоя бетона соблюдена ?",
    "Арматурные каркасы закреплены ?",
    "Опалубка обеспечивает проектные размеры ?",
    "Бетон уплотнён вибрированием ?",
]

# ---------------- Логирование ----------------
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO
)
logger = logging.getLogger(__name__)


# ---------------- Telegram API ----------------
class TelegramBot:
    def __init__(self, token):
        self.token = token
        self.base_url = f"https://api.telegram.org/bot{token}"
        self.offset = 0

    def get_updates(self):
        try:
            response = requests.get(f"{self.base_url}/getUpdates", params={"offset": self.offset, "timeout": 30})
            return response.json().get("result", [])
        except Exception as e:
            logger.error(f"Error getting updates: {e}")
            return []

    def send_message(self, chat_id, text, reply_markup=None, parse_mode=None):
        data = {
            "chat_id": chat_id,
            "text": text
        }
        if reply_markup:
            data["reply_markup"] = json.dumps(reply_markup)
        if parse_mode:
            data["parse_mode"] = parse_mode

        requests.post(f"{self.base_url}/sendMessage", json=data)

    def edit_message(self, chat_id, message_id, text, reply_markup=None):
        data = {
            "chat_id": chat_id,
            "message_id": message_id,
            "text": text
        }
        if reply_markup:
            data["reply_markup"] = json.dumps(reply_markup)

        requests.post(f"{self.base_url}/editMessageText", json=data)

    def send_document(self, chat_id, document_path, caption=""):
        with open(document_path, 'rb') as file:
            files = {'document': file}
            data = {'chat_id': chat_id, 'caption': caption}
            requests.post(f"{self.base_url}/sendDocument", files=files, data=data)


# ---------------- Работа с базой ----------------
def init_db():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS inspections (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            chat_id INTEGER,
            created_at TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            inspection_id INTEGER,
            idx INTEGER,
            text TEXT,
            answer TEXT,
            comment TEXT,
            photo_path TEXT
        )
    """)
    conn.commit()
    conn.close()


def create_inspection(chat_id):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    now = datetime.utcnow().isoformat()
    cur.execute("INSERT INTO inspections (chat_id, created_at) VALUES (?, ?)", (chat_id, now))
    ins_id = cur.lastrowid
    for i, text in enumerate(CHECKLIST):
        cur.execute("INSERT INTO items (inspection_id, idx, text) VALUES (?, ?, ?)", (ins_id, i, text))
    conn.commit()
    conn.close()
    return ins_id


def update_item(inspection_id, idx, **kwargs):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    for k, v in kwargs.items():
        cur.execute(f"UPDATE items SET {k}=? WHERE inspection_id=? AND idx=?", (v, inspection_id, idx))
    conn.commit()
    conn.close()


def get_items(inspection_id):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT idx, text, answer, comment, photo_path FROM items WHERE inspection_id=? ORDER BY idx",
                (inspection_id,))
    rows = cur.fetchall()
    conn.close()
    return rows


def is_inspection_complete(inspection_id):
    items = get_items(inspection_id)
    for idx, text, answer, comment, photo_path in items:
        if answer is None:
            return False, f"Не заполнен пункт {idx + 1}"
    return True, None


# ---------------- Загрузка шаблона ----------------
def load_template():
    if os.path.exists(TEMPLATE_PATH):
        return Document(TEMPLATE_PATH)
    else:
        logger.warning(f"Шаблон {TEMPLATE_PATH} не найден, используется пустой документ")
        doc = Document()
        doc.add_heading('Акт обследования фундамента', level=1)
        return doc


# ---------------- Функции для работы с таблицами в Word ----------------
def find_table_marker(paragraphs, marker_text):
    """Находит параграф с маркером для вставки таблицы"""
    for i, paragraph in enumerate(paragraphs):
        if marker_text in paragraph.text:
            return i, paragraph
    return None, None


def insert_table_at_marker(doc, items, marker_text="[TABLE_PLACEHOLDER]"):
    """Вставляет таблицу вместо текстового маркера"""

    # Ищем маркер в параграфах
    marker_index, marker_paragraph = find_table_marker(doc.paragraphs, marker_text)

    if marker_paragraph is None:
        # Если маркер не найден, добавляем таблицу в конец
        logger.warning(f"Маркер '{marker_text}' не найден, добавляем таблицу в конец")
        return add_table_to_end(doc, items)

    # Создаем таблицу
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№"
    hdr_cells[1].text = "Пункт проверки"
    hdr_cells[2].text = "Соответствует"
    hdr_cells[3].text = "Комментарий"

    # Данные таблицы
    for idx, text, answer, comment, photo_path in items:
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[1].text = text
        row_cells[2].text = answer or ""
        row_cells[3].text = comment or ""

    # Получаем XML элемент таблицы
    tbl_element = table._tbl

    # Заменяем параграф с маркером на таблицу
    marker_paragraph._p.getparent().replace(marker_paragraph._p, tbl_element)

    return doc


def add_table_to_end(doc, items):
    """Добавляет таблицу в конец документа (запасной вариант)"""
    doc.add_paragraph("Результаты проверки:")

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№"
    hdr_cells[1].text = "Пункт проверки"
    hdr_cells[2].text = "Соответствует"
    hdr_cells[3].text = "Комментарий"

    for idx, text, answer, comment, photo_path in items:
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[1].text = text
        row_cells[2].text = answer or ""
        row_cells[3].text = comment or ""

    return doc


# ---------------- Клавиатуры ----------------
def get_welcome_keyboard():
    return {
        "inline_keyboard": [
            [
                {"text": "🚀 Начать проверку", "callback_data": "start_inline"},
                {"text": "ℹ️ О боте", "callback_data": "about_bot"}
            ],
            [
                {"text": "📖 Инструкция", "callback_data": "help_inline"}
            ]
        ]
    }


def get_main_menu_keyboard(ins_id):
    items = get_items(ins_id)
    keyboard = []

    for i, item in enumerate(CHECKLIST):
        status = "✅" if items[i][2] == "Да" else "❌" if items[i][2] == "Нет" else "⏳"
        keyboard.append([{
            "text": f"{status} {i + 1}. {item}",
            "callback_data": f"item:{i}"
        }])

    keyboard.append([
        {"text": "📊 Генерировать отчёт", "callback_data": "generate"}
    ])

    keyboard.append([
        {"text": "ℹ️ О боте", "callback_data": "about_bot"}
    ])

    return {"inline_keyboard": keyboard}


def get_item_keyboard(idx):
    return {
        "inline_keyboard": [
            [
                {"text": "✅ Да", "callback_data": f"set:{idx}:Да"},
                {"text": "❌ Нет", "callback_data": f"set:{idx}:Нет"}
            ],
            [
                {"text": "📝 Добавить комментарий", "callback_data": f"comment:{idx}"}
            ],
            [
                {"text": "🔙 Назад к списку", "callback_data": "back"}
            ]
        ]
    }


# ---------------- Обработчики ----------------
class BotHandler:
    def __init__(self, bot):
        self.bot = bot
        self.user_states = {}  # {chat_id: {"state": "comment", "idx": 0, "ins_id": 1}}

    def handle_update(self, update):
        if "message" in update:
            self.handle_message(update["message"])
        elif "callback_query" in update:
            self.handle_callback(update["callback_query"])

    def handle_message(self, message):
        chat_id = message["chat"]["id"]
        text = message.get("text", "").strip()

        # Обработка команд
        if text == "/start":
            self.handle_start(chat_id)
        elif text == "/help":
            self.handle_help(chat_id)
        elif text == "/about":
            self.handle_about(chat_id)
        else:
            # Обработка состояний
            if chat_id in self.user_states:
                state = self.user_states[chat_id]
                if state["state"] == "comment":
                    idx = state["idx"]
                    ins_id = state["ins_id"]
                    update_item(ins_id, idx, comment=text)
                    self.bot.send_message(chat_id, f"✅ Комментарий для пункта {idx + 1} сохранён.")
                    self.show_main_menu(chat_id, ins_id)
                    del self.user_states[chat_id]
                    return
            # Если пользователь просто написал что-то без команды - отправляем приветствие
            self.send_welcome_message(chat_id)

    def handle_callback(self, callback_query):
        chat_id = callback_query["message"]["chat"]["id"]
        message_id = callback_query["message"]["message_id"]
        data = callback_query["data"]

        if data == "start_inline":
            self.handle_start(chat_id)
            return
        elif data == "about_bot":
            self.handle_about(chat_id)
            return
        elif data == "help_inline":
            self.handle_help(chat_id)
            return

        ins_id = self.get_inspection_id(chat_id)

        if data.startswith("item:"):
            idx = int(data.split(":")[1])
            self.show_item_menu(chat_id, message_id, idx, ins_id)

        elif data.startswith("set:"):
            _, idx_str, ans = data.split(":")
            idx = int(idx_str)
            update_item(ins_id, idx, answer=ans)
            self.bot.edit_message(chat_id, message_id, f"✅ Ответ для пункта {idx + 1} установлен: {ans}")
            self.show_main_menu(chat_id, ins_id)

        elif data.startswith("comment:"):
            idx = int(data.split(":")[1])
            self.user_states[chat_id] = {"state": "comment", "idx": idx, "ins_id": ins_id}
            self.bot.edit_message(chat_id, message_id, "✍️ Отправьте текст комментария для этого пункта.")

        elif data == "generate":
            self.generate_report(chat_id, message_id, ins_id)

        elif data == "back":
            self.show_main_menu(chat_id, ins_id)

    def send_welcome_message(self, chat_id):
        welcome_text = """
🏗️ *Бот для обследования фундамента*

*Кратко о возможностях:*

✅ *Проверка фундамента* по 5 ключевым параметрам
📝 *Добавление комментариев* к каждому пункту  
📊 *Генерация актов* в формате Word
💾 *Сохранение истории* проверок

*Бот готов к работе! Выберите действие ниже:*
        """
        self.bot.send_message(chat_id, welcome_text,
                              reply_markup=get_welcome_keyboard(),
                              parse_mode="Markdown")

    def handle_start(self, chat_id):
        welcome_text = """
👋 *Начинаем проверку фундамента!*

Заполните все 5 пунктов проверки:
1. Марка бетона
2. Защитный слой бетона  
3. Арматурные каркасы
4. Опалубка
5. Уплотнение бетона

*После заполнения всех пунктов вы получите готовый акт в Word формате!*
        """

        ins_id = create_inspection(chat_id)
        self.bot.send_message(chat_id, welcome_text, parse_mode="Markdown")
        self.show_main_menu(chat_id, ins_id)

    def handle_about(self, chat_id):
        about_text = """🏗️ *Бот для проведения аудита*

*Что умеет этот бот:*

📋 *Проведение проверок* 
- Систематизированный чек-лист по всем ключевым параметрам фундамента
- Поэтапное заполнение каждого пункта

✅ *Интуитивный интерфейс*
- Встроенные клавиатуры для быстрых ответов
- Статусы выполнения в реальном времени
- Возможность добавления комментариев и фото

📊 *Автоматическая отчетность*
- Генерация профессиональных актов в формате Word
- Структурированные таблицы с результатами проверки
- Готовые шаблоны документов

💾 *Надежное хранение*
- Все данные сохраняются в базе
- История проведенных проверок
- Возможность возобновить незавершенные проверки

*Основные команды:*
/start - начать новую проверку
/help - помощь по использованию
/about - информация о боте

*Для начала работы просто нажмите "Начать проверку"*"""
        self.bot.send_message(chat_id, about_text, parse_mode="Markdown")

    def handle_help(self, chat_id):
        help_text = """
🆘 *Помощь по использованию бота*

*Как работать с ботом:*

1. *Начало работы* 
   - Нажмите /start для создания новой проверки
   - Или выберите пункт из активного меню

2. *Заполнение пунктов*
   - Нажмите на любой пункт для его заполнения
   - Выберите ✅ Да или ❌ Нет
   - Добавьте комментарий если нужно

3. *Дополнительные возможности*
   - Добавление фото к пунктам
   - Редактирование ранее заполненных данных
   - Просмотр прогресса заполнения

4. *Генерация отчета*
   - Когда все пункты заполнены - нажмите "Генерировать отчёт"
   - Получите готовый Word-документ

*Статусы пунктов:*
✅ - соответствует требованиям
❌ - не соответствует требованиям  
⏳ - не заполнено

*Команды:*
/start - начать проверку
/help - эта справка  
/about - информация о боте

*Если что-то не работает:*
- Проверьте подключение к интернету
- Убедитесь, что все пункты заполнены перед генерацией отчета
- При проблемах - перезапустите бота командой /start
        """
        self.bot.send_message(chat_id, help_text, parse_mode="Markdown")

    def show_main_menu(self, chat_id, ins_id):
        self.bot.send_message(
            chat_id,
            "Начинаем обследование. Выберите пункт:\n✅ - соответствует, ❌ - не соответствует, ⏳ - не заполнено",
            get_main_menu_keyboard(ins_id)
        )

    def show_item_menu(self, chat_id, message_id, idx, ins_id):
        items = get_items(ins_id)
        current_item = items[idx]

        status_info = ""
        if current_item[2]:
            status_info = f"\n\nТекущий статус: {current_item[2]}"
        if current_item[3]:
            status_info += f"\nКомментарий: {current_item[3]}"

        self.bot.edit_message(
            chat_id,
            message_id,
            f"Пункт {idx + 1}: {CHECKLIST[idx]}{status_info}",
            get_item_keyboard(idx)
        )

    def generate_report(self, chat_id, message_id, ins_id):
        is_complete, reason = is_inspection_complete(ins_id)
        if not is_complete:
            self.bot.edit_message(chat_id, message_id,
                                  f"❌ Не все пункты заполнены: {reason}\nЗаполните все пункты перед генерацией отчёта.")
            return

        self.bot.edit_message(chat_id, message_id, "⏳ Генерирую отчёт...")

        try:
            items = get_items(ins_id)
            doc = load_template()

            # Вставляем таблицу в маркер или в конец
            doc = insert_table_at_marker(doc, items)

            with tempfile.NamedTemporaryFile(prefix="report_", suffix=".docx", delete=False) as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name

            self.bot.send_document(chat_id, tmp_path, f"📊 Отчёт по проверке #{ins_id}")
            os.unlink(tmp_path)

        except Exception as e:
            logger.error(f"Ошибка генерации отчёта: {e}")
            self.bot.send_message(chat_id, "❌ Произошла ошибка при генерации отчёта.")

    def get_inspection_id(self, chat_id):
        conn = sqlite3.connect(DB_PATH)
        cur = conn.cursor()
        cur.execute("SELECT id FROM inspections WHERE chat_id=? ORDER BY id DESC LIMIT 1", (chat_id,))
        result = cur.fetchone()
        conn.close()

        if not result:
            return create_inspection(chat_id)
        return result[0]


# ---------------- Главная функция ----------------
def main():
    # Создаем директории
    os.makedirs(os.path.join(os.path.dirname(__file__), "data", "photos"), exist_ok=True)

    init_db()
    bot = TelegramBot(BOT_TOKEN)
    handler = BotHandler(bot)

    logger.info("Bot started - Press Ctrl+C to stop")

    try:
        while True:
            updates = bot.get_updates()
            for update in updates:
                bot.offset = update["update_id"] + 1
                handler.handle_update(update)
    except KeyboardInterrupt:
        logger.info("Bot stopped by user")


if __name__ == "__main__":
    main()